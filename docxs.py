"""
Creates a report in the form of a docx file for a minion run
"""

import docx

def docxs(selection_barcode,date, flowcell_id, barcode_present, result_directory,design_file_directory = ''):
    report_writing_directory = result_directory
    print(report_writing_directory)

    ##### Watch out change in the directory#####
    doc = docx.Document()
    doc.add_heading('Rapport run Minion\n', level=1)
    doc.add_paragraph('Flowcell id: '+flowcell_id)
    doc.add_paragraph('Date: '+date)
    doc.add_paragraph('Histogram of read counts according to the type read')
    doc.add_picture(report_writing_directory+'images/read_count_histogram.png', width=docx.shared.Inches(6),height=docx.shared.Cm(9))
    doc.add_paragraph('Phred score according to the read type')
    doc.add_picture(report_writing_directory+'images/read_quality_boxplot.png', width=docx.shared.Inches(5),height=docx.shared.Cm(9))
    doc.add_paragraph('Channel counts')
    doc.add_picture(report_writing_directory+'images/channel_count_histogram.png', width=docx.shared.Inches(5), height=docx.shared.Cm(9))
    doc.add_paragraph('Curve representing the reads number produced against the time')
    doc.add_picture(report_writing_directory+'images/read_number_run.png', width=docx.shared.Inches(5), height=docx.shared.Cm(9))
    doc.add_paragraph('Channel occupancy')
    doc.add_picture(report_writing_directory+'images/channel_occupancy.png', width=docx.shared.Inches(8), height=docx.shared.Cm(11))
    doc.add_paragraph('Read size histogram')
    doc.add_picture(report_writing_directory+'images/read_length_histogram.png', width=docx.shared.Inches(8), height=docx.shared.Cm(11))
    doc.add_paragraph('Phred score frequency')
    doc.add_picture(report_writing_directory+'images/phred_score_frequency.png', width=docx.shared.Inches(5), height=docx.shared.Cm(9))
    doc.add_paragraph("Relation between the sequence length template and the mean qscore template")
    doc.add_picture(report_writing_directory+'images/scatter_plot.png', width=docx.shared.Inches(5), height=docx.shared.Cm(9))
    if barcode_present:
        doc.add_paragraph("Barcode pie chart")
        doc.add_picture(report_writing_directory+'images/barcode_percentage_pie_chart.png', width=docx.shared.Inches(5), height=docx.shared.Cm(9))
        selection_barcode = selection_barcode[:-1]
        doc.add_paragraph('Read length boxplots')
        doc.add_picture(report_writing_directory+'images/barcode_total.png',width=docx.shared.Inches(5), height=docx.shared.Cm(9))
        doc.add_paragraph('Boxplots of phred score distribution by barcode')
        doc.add_picture(report_writing_directory+"images/barcode_phred_score_boxplot.png",width=docx.shared.Inches(5), height=docx.shared.Cm(9))

       # for barcode_image in selection_barcode:
        #    doc.add_picture('images/image_{}.png'.format(barcode_image), width=docx.shared.Inches(5), height=docx.shared.Cm(9))

        doc.paragraphs[0].runs[0].underline = True
        doc.paragraphs[3].runs[0].underline = True
        doc.paragraphs[5].runs[0].underline = True
        doc.paragraphs[7].runs[0].underline = True
        doc.paragraphs[9].runs[0].underline = True
        doc.paragraphs[11].runs[0].underline = True
        doc.paragraphs[13].runs[0].underline = True
        doc.paragraphs[15].runs[0].underline = True
        doc.paragraphs[17].runs[0].underline = True
        doc.paragraphs[19].runs[0].underline = True
        doc.paragraphs[21].runs[0].underline = True
    doc.save(result_directory+'Rapport.docx')

